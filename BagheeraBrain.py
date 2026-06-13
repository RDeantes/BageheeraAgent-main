import os
import shutil
import subprocess
from datetime import datetime
from docx import Document

from BagheeraExcel import contrato_desde_excel, agregar_empleado
from vigencias import listar_vencidos


# 🧠 MEMORIA GLOBAL
estado_contrato = {}
estado_empleado = {}
estado_renovacion = {}


# =========================================================
# 🐱 PERSONALIDAD
# =========================================================
def personalidad_bagheera(respuesta):
    return f"🐱:\n{respuesta}\n\n“La fuerza no está en rugir… está en saber cuándo moverte en silencio.” 🌑"


# =========================================================
# 🔤 NORMALIZAR TEXTO (PRO)
# =========================================================
def normalizar(texto):
    texto = texto.lower().strip()
    reemplazos = {
        "á": "a", "é": "e", "í": "i",
        "ó": "o", "ú": "u"
    }
    for k, v in reemplazos.items():
        texto = texto.replace(k, v)
    return texto


# =========================================================
# 🎯 ROUTER PRINCIPAL
# =========================================================
def procesar(mensaje):
    global estado_contrato, estado_empleado, estado_renovacion

    mensaje = normalizar(mensaje)

    if estado_renovacion.get("activo"):
        return flujo_renovacion(mensaje)

    # 🔥 PRIORIDAD: EMPLEADO
    if estado_empleado.get("activo"):
        return flujo_agregar_empleado(mensaje)

    # 🔥 DESPUÉS CONTRATO
    if estado_contrato.get("activo"):
        return flujo_contrato(mensaje)

    # =====================================================
    # COMANDOS
    # =====================================================

    # 🌴 VACACIONES (tolerante)
    if "vacacion" in mensaje or "vacac" in mensaje:
        return revisar_vacaciones_por_mes(mensaje)

    # 🔎 VIGENCIAS / RENOVACIONES
    if "vigencia" in mensaje or "vigencias" in mensaje:
        return revisar_vigencias()

    if mensaje.startswith("renovar"):
        nombre = mensaje.replace("renovar", "", 1).strip()
        if not nombre:
            return personalidad_bagheera("Indica el nombre del empleado para renovar, por ejemplo: RENOVAR JUAN PEREZ")
        estado_renovacion = {"activo": True, "nombre": nombre.upper()}
        return personalidad_bagheera("Tipo de contrato? (temporal / permanente)")

    # 👤 AGREGAR EMPLEADO
    if "agregar empleado" in mensaje:
        estado_contrato = {}
        estado_empleado = {"activo": True, "paso": 1}
        return personalidad_bagheera("Nombre del empleado:")

    # 📄 NUEVO CONTRATO
    if "nuevo contrato" in mensaje:
        estado_empleado = {}
        estado_contrato = {"activo": True}
        return personalidad_bagheera("¿Tipo de contrato? (temporal / permanente)")

    return personalidad_bagheera("No entendí la orden.")


# =========================================================
# 🧠 FLUJO CONTRATO
# =========================================================
def flujo_contrato(mensaje):
    global estado_contrato

    if "tipo" not in estado_contrato:
        if "temporal" in mensaje:
            estado_contrato["tipo"] = "TEMPORAL"
        elif "permanente" in mensaje or "indeterminado" in mensaje:
            estado_contrato["tipo"] = "INDETERMINADO"
        else:
            return personalidad_bagheera("Responde: temporal o permanente")
        return personalidad_bagheera("¿Jornada completa o parcial?")

    elif "jornada" not in estado_contrato:
        if "completa" in mensaje:
            estado_contrato["jornada"] = "COMPLETA"
            return personalidad_bagheera("Duración del contrato:")
        elif "parcial" in mensaje:
            estado_contrato["jornada"] = "PARCIAL"
            return personalidad_bagheera("¿Qué días trabajará? (ej: LUNES, MIERCOLES Y VIERNES):")
        else:
            return personalidad_bagheera("Responde: completa o parcial")

    elif estado_contrato.get("jornada") == "PARCIAL" and "dias" not in estado_contrato:
        estado_contrato["dias"] = mensaje.upper()
        return personalidad_bagheera("Duración del contrato:")

    elif "duracion" not in estado_contrato:
        estado_contrato["duracion"] = mensaje.upper()
        return personalidad_bagheera("Fecha inicio (YYYY-MM-DD):")

    elif "fecha_inicio" not in estado_contrato:
        estado_contrato["fecha_inicio"] = mensaje
        return personalidad_bagheera("Fecha fin (YYYY-MM-DD):")

    elif "fecha_termino" not in estado_contrato:
        estado_contrato["fecha_termino"] = mensaje
        return personalidad_bagheera("Nombre del empleado:")

    elif "nombre" not in estado_contrato:
        estado_contrato["nombre"] = mensaje.upper()

        datos = estado_contrato.copy()
        estado_contrato = {}

        return contrato_desde_excel(datos, generar_contrato, personalidad_bagheera)


# =========================================================
# 👤 FLUJO AGREGAR EMPLEADO
# =========================================================
def flujo_renovacion(mensaje):
    global estado_renovacion

    if "tipo" not in estado_renovacion:
        if "temporal" in mensaje:
            estado_renovacion["tipo"] = "TEMPORAL"
        elif "permanente" in mensaje or "indeterminado" in mensaje:
            estado_renovacion["tipo"] = "INDETERMINADO"
        else:
            return personalidad_bagheera("Responde: temporal o permanente")
        return personalidad_bagheera("¿Jornada completa o parcial?")

    if "jornada" not in estado_renovacion:
        if "completa" in mensaje:
            estado_renovacion["jornada"] = "COMPLETA"
            return personalidad_bagheera("Duración del contrato:")
        elif "parcial" in mensaje:
            estado_renovacion["jornada"] = "PARCIAL"
            return personalidad_bagheera("¿Qué días trabajará? (ej: LUNES, MIERCOLES Y VIERNES):")
        else:
            return personalidad_bagheera("Responde: completa o parcial")

    if estado_renovacion.get("jornada") == "PARCIAL" and "dias" not in estado_renovacion:
        estado_renovacion["dias"] = mensaje.upper()
        return personalidad_bagheera("Duración del contrato:")

    if "duracion" not in estado_renovacion:
        estado_renovacion["duracion"] = mensaje.upper()
        return personalidad_bagheera("Fecha inicio (YYYY-MM-DD):")

    if "fecha_inicio" not in estado_renovacion:
        estado_renovacion["fecha_inicio"] = mensaje
        return personalidad_bagheera("Fecha fin (YYYY-MM-DD):")

    if "fecha_termino" not in estado_renovacion:
        estado_renovacion["fecha_termino"] = mensaje
        datos = {
            "nombre": estado_renovacion["nombre"],
            "tipo": estado_renovacion["tipo"],
            "jornada": estado_renovacion["jornada"],
            "dias": estado_renovacion.get("dias", "LUNES A SABADO"),
            "duracion": estado_renovacion["duracion"],
            "fecha_inicio": estado_renovacion["fecha_inicio"],
            "fecha_termino": estado_renovacion["fecha_termino"],
        }
        estado_renovacion = {}
        return contrato_desde_excel(datos, generar_contrato, personalidad_bagheera)


def revisar_vigencias():
    vencidos = listar_vencidos()
    if not vencidos:
        return personalidad_bagheera("✅ No hay contratos vencidos en la fuente actual.")

    texto = "⚠️ Contratos vencidos:\n\n" + "\n".join(
        f"- {item['nombre']} (vigencia: {item['vigencia']})"
        for item in vencidos
    )
    texto += "\n\nResponde: RENOVAR NOMBRE para generar un nuevo contrato."
    return personalidad_bagheera(texto)


def flujo_agregar_empleado(mensaje):
    global estado_empleado

    paso = estado_empleado["paso"]

    if paso == 1:
        estado_empleado["NOMBRE"] = mensaje.upper()
        estado_empleado["paso"] = 2
        return personalidad_bagheera("Área:")

    elif paso == 2:
        estado_empleado["AREA"] = mensaje.upper()
        estado_empleado["paso"] = 3
        return personalidad_bagheera("Puesto:")

    elif paso == 3:
        estado_empleado["PUESTO"] = mensaje.upper()
        estado_empleado["paso"] = 4
        return personalidad_bagheera("Fecha ingreso (YYYY-MM-DD):")

    elif paso == 4:
        estado_empleado["FECHA_INGRESO"] = mensaje
        estado_empleado["paso"] = 5
        return personalidad_bagheera("Tipo contrato:")

    elif paso == 5:
        estado_empleado["TIPO_CONTRATO"] = mensaje.upper()
        estado_empleado["paso"] = 6
        return personalidad_bagheera("Vigencia:")

    elif paso == 6:
        estado_empleado["VIGENCIA"] = mensaje
        estado_empleado["paso"] = 7
        return personalidad_bagheera("Nacionalidad:")

    elif paso == 7:
        estado_empleado["NACIONALIDAD"] = mensaje.upper()
        estado_empleado["paso"] = 8
        return personalidad_bagheera("Sexo:")

    elif paso == 8:
        estado_empleado["SEXO"] = mensaje.upper()
        estado_empleado["paso"] = 9
        return personalidad_bagheera("CURP:")

    elif paso == 9:
        estado_empleado["CURP"] = mensaje.upper()
        estado_empleado["paso"] = 10
        return personalidad_bagheera("Domicilio:")

    elif paso == 10:
        estado_empleado["DOMICILIO"] = mensaje.upper()

        agregar_empleado(estado_empleado)

        estado_empleado = {}

        return personalidad_bagheera("✅ Empleado agregado correctamente")


# =========================================================
# 📄 GENERAR CONTRATO
# =========================================================
def _reemplazar_placeholders_en_parrafos(elemento, datos):
    for p in elemento.paragraphs:
        if not p.text:
            continue
        texto = p.text
        for k, v in datos.items():
            texto = texto.replace(f"{{{{{k}}}}}", str(v))
        if texto != p.text:
            p.text = texto

    for tabla in elemento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                _reemplazar_placeholders_en_parrafos(celda, datos)


def _seleccionar_plantilla(datos):
    tipo = str(datos.get("tipo", "")).upper()
    jornada = str(datos.get("jornada", "")).upper()

    if "INDETERMINADO" in tipo or "PERMANENTE" in tipo:
        if "PARCIAL" in jornada:
            return "CONTRATO FORMATO TIEMPO INDETERMINADO Y PARCIAL.docx"
        return "CONTRATO FORMATO TIEMPO INDETERMINADO.docx"
    else:
        if "PARCIAL" in jornada:
            return "CONTRATO FORMATO TEMPORAL Y PARCIAL.docx"
        return "CONTRATO FORMATO TEMPORAL.docx"


def generar_contrato(datos):
    base_path = os.path.dirname(__file__)

    plantilla = os.path.join(base_path, _seleccionar_plantilla(datos))

    doc = Document(plantilla)

    _reemplazar_placeholders_en_parrafos(doc, datos)

    nombre_sanitizado = "CONTRATO" if not datos.get("nombre") else str(datos.get("nombre")).strip()
    nombre_sanitizado = "".join(ch if ch.isalnum() or ch in " _-." else "_" for ch in nombre_sanitizado)
    nombre_sanitizado = nombre_sanitizado.strip() or "CONTRATO"

    ruta_docx = os.path.join(base_path, f"{nombre_sanitizado}.docx")
    doc.save(ruta_docx)

    ruta_pdf = ruta_docx.replace(".docx", ".pdf")

    try:
        if shutil.which("soffice"):
            subprocess.run([
                "soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", base_path,
                ruta_docx
            ], check=True)
        else:
            from docx2pdf import convert
            convert(ruta_docx, ruta_pdf)
    except Exception:
        if not os.path.exists(ruta_pdf):
            return ruta_docx

    if not os.path.exists(ruta_pdf) and os.path.exists(ruta_docx):
        return ruta_docx

    return ruta_pdf


def revisar_vacaciones_por_mes(mensaje):
    from google_sheets import get_sheet
    from datetime import datetime, timedelta

    meses = {
        "enero": 1, "febrero": 2, "marzo": 3,
        "abril": 4, "mayo": 5, "junio": 6,
        "julio": 7, "agosto": 8, "septiembre": 9,
        "octubre": 10, "noviembre": 11, "diciembre": 12
    }

    mes_detectado = None

    for mes, num in meses.items():
        if mes in mensaje:
            mes_detectado = num
            break

    if not mes_detectado:
        return personalidad_bagheera("❌ Mes no válido")

    try:
        sheet = get_sheet()
        registros = sheet.get_all_records()
    except Exception as exc:
        return personalidad_bagheera(f"⚠️ No se puede consultar Google Sheets ahora: {exc}")

    resultado = []

    # 🔥 función robusta de parseo
    def parse_fecha(valor):
        if not valor:
            return None

        # caso datetime real
        if isinstance(valor, datetime):
            return valor

        texto = str(valor).strip()

        # 🔥 limpiar formato tipo ISO
        if "T" in texto:
            texto = texto.split("T")[0]

        # intentar formatos comunes
        formatos = [
            "%Y-%m-%d",
            "%d/%m/%Y",
            "%m/%d/%Y",
            "%d-%m-%Y"
        ]

        for f in formatos:
            try:
                return datetime.strptime(texto, f)
            except:
                continue

        # 🔥 caso número (Excel serial date)
        try:
            num = float(texto)
            base = datetime(1899, 12, 30)
            return base + timedelta(days=num)
        except:
            pass

        return None

    # 🔍 procesamiento real
    for row in registros:
        fecha_raw = row.get("FECHA_DE_INGRESO", "")
        fecha = parse_fecha(fecha_raw)

        if fecha:
            if fecha.month == mes_detectado:
                resultado.append(row.get("NOMBRE", ""))

    if not resultado:
        return personalidad_bagheera("No hay empleados con aniversario en ese mes")

    return personalidad_bagheera(
        "Empleados con aniversario:\n\n" + "\n".join(resultado)
    )
    
def generar_contrato_directo(nombre, fecha_inicio, fecha_termino):

    datos = {
        "tipo": "TEMPORAL",  # puedes hacerlo dinámico luego
        "jornada": "COMPLETA",
        "duracion": "RENOVACIÓN",
        "fecha_inicio": fecha_inicio,
        "fecha_termino": fecha_termino,
        "nombre": nombre
    }

    return contrato_desde_excel(
        datos,
        generar_contrato,
        personalidad_bagheera
        
        
        
    )